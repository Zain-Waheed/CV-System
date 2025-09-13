from flask import Flask, request, jsonify
from flask_cors import CORS
import PyPDF2
import docx
import re
import os
from werkzeug.utils import secure_filename
from collections import Counter
import json
from datetime import datetime, timedelta

app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
UPLOAD_FOLDER = 'temp_uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

# Create upload folder if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class EnhancedCVAnalyzer:
    def __init__(self):
        # Multi-industry skill keywords
        self.skill_keywords = {
            # Technology & IT
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin', 'typescript', 'scala', 'r', 'matlab', 'perl'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'bootstrap', 'jquery', 'sass', 'webpack', 'next.js', 'gatsby'],
            'data_science': ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'data analysis', 'statistics', 'ml', 'ai', 'tableau', 'power bi'],
            'databases': ['sql', 'mysql', 'mongodb', 'postgresql', 'redis', 'elasticsearch', 'oracle', 'sqlite', 'nosql', 'cassandra', 'dynamodb'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'ci/cd', 'devops', 'serverless', 'microservices'],
            'mobile': ['android', 'ios', 'react native', 'flutter', 'swift', 'kotlin', 'xamarin', 'cordova', 'ionic'],
            'cybersecurity': ['penetration testing', 'ethical hacking', 'cissp', 'cism', 'firewall', 'encryption', 'vulnerability assessment', 'incident response'],
            
            # Healthcare & Medical
            'medical_skills': ['patient care', 'medical diagnosis', 'surgery', 'emergency medicine', 'radiology', 'anesthesia', 'cardiology', 'neurology', 'oncology', 'pediatrics'],
            'healthcare_tech': ['emr', 'ehr', 'epic', 'cerner', 'meditech', 'medical imaging', 'telemedicine', 'healthcare analytics'],
            'nursing': ['critical care', 'patient assessment', 'medication administration', 'wound care', 'iv therapy', 'patient education', 'bls', 'acls', 'pals'],
            'pharmacy': ['pharmaceutical care', 'drug interactions', 'compounding', 'clinical pharmacy', 'pharmacokinetics', 'medication therapy management'],
            
            # Finance & Accounting
            'finance': ['financial analysis', 'investment banking', 'portfolio management', 'risk management', 'derivatives', 'equity research', 'forex', 'financial modeling'],
            'accounting': ['gaap', 'ifrs', 'tax preparation', 'auditing', 'bookkeeping', 'financial reporting', 'cost accounting', 'forensic accounting'],
            'fintech': ['blockchain', 'cryptocurrency', 'robo-advisor', 'algorithmic trading', 'payment processing', 'regulatory compliance'],
            
            # Marketing & Sales
            'digital_marketing': ['seo', 'sem', 'social media marketing', 'content marketing', 'email marketing', 'ppc', 'google analytics', 'facebook ads', 'conversion optimization'],
            'traditional_marketing': ['brand management', 'market research', 'advertising', 'public relations', 'event marketing', 'print media', 'trade shows'],
            'sales': ['lead generation', 'crm', 'salesforce', 'hubspot', 'b2b sales', 'b2c sales', 'account management', 'sales forecasting', 'negotiation'],
            
            # Education & Training
            'teaching': ['curriculum development', 'lesson planning', 'classroom management', 'educational technology', 'assessment', 'differentiated instruction'],
            'educational_tech': ['lms', 'moodle', 'blackboard', 'canvas', 'e-learning', 'instructional design', 'educational apps'],
            
            # Engineering & Manufacturing
            'mechanical_engineering': ['cad', 'solidworks', 'autocad', 'finite element analysis', 'thermodynamics', 'fluid mechanics', 'manufacturing processes'],
            'electrical_engineering': ['circuit design', 'power systems', 'control systems', 'plc programming', 'embedded systems', 'signal processing'],
            'civil_engineering': ['structural design', 'construction management', 'surveying', 'geotechnical engineering', 'transportation engineering'],
            'manufacturing': ['lean manufacturing', 'six sigma', 'quality control', 'supply chain management', 'production planning', 'process improvement'],
            
            # Creative & Design
            'graphic_design': ['photoshop', 'illustrator', 'indesign', 'figma', 'sketch', 'after effects', 'premiere pro', 'typography', 'branding'],
            'web_design': ['ui design', 'ux design', 'wireframing', 'prototyping', 'user research', 'responsive design', 'accessibility'],
            'multimedia': ['video editing', 'motion graphics', 'animation', '3d modeling', 'blender', 'maya', 'cinema 4d'],
            
            # Legal & Compliance
            'legal': ['contract law', 'litigation', 'legal research', 'compliance', 'intellectual property', 'corporate law', 'family law', 'criminal law'],
            'regulatory': ['gdpr', 'hipaa', 'sox', 'regulatory compliance', 'risk assessment', 'policy development'],
            
            # Human Resources
            'hr': ['recruitment', 'talent acquisition', 'performance management', 'employee relations', 'compensation', 'benefits administration', 'hris'],
            'hr_tech': ['workday', 'successfactors', 'bamboohr', 'adp', 'applicant tracking system', 'payroll systems'],
            
            # Operations & Supply Chain
            'operations': ['operations management', 'process optimization', 'logistics', 'inventory management', 'vendor management', 'procurement'],
            'supply_chain': ['supply chain optimization', 'demand planning', 'erp', 'sap', 'oracle', 'warehouse management'],
            
            # Language & Communication
            'languages': ['english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'arabic', 'hindi', 'portuguese', 'russian'],
            'communication': ['technical writing', 'copywriting', 'presentation skills', 'public speaking', 'stakeholder management', 'cross-cultural communication'],
            
            # Soft Skills (Enhanced)
            'leadership': ['team leadership', 'strategic planning', 'change management', 'coaching', 'mentoring', 'conflict resolution', 'decision making'],
            'project_management': ['agile', 'scrum', 'kanban', 'pmp', 'prince2', 'waterfall', 'risk management', 'stakeholder management'],
            'analytical': ['problem solving', 'critical thinking', 'data analysis', 'research', 'troubleshooting', 'root cause analysis'],
            'interpersonal': ['teamwork', 'collaboration', 'customer service', 'relationship building', 'emotional intelligence', 'active listening']
        }
        
        # Industry-specific keywords for automatic CV type detection
        self.industry_keywords = {
            'technology': ['software', 'developer', 'engineer', 'programming', 'coding', 'algorithm', 'database', 'system', 'technical', 'it'],
            'healthcare': ['medical', 'nurse', 'doctor', 'physician', 'patient', 'clinical', 'hospital', 'healthcare', 'medical'],
            'finance': ['financial', 'accounting', 'investment', 'banking', 'finance', 'audit', 'tax', 'budget', 'revenue'],
            'marketing': ['marketing', 'sales', 'advertising', 'brand', 'campaign', 'customer', 'market', 'promotion'],
            'education': ['teaching', 'education', 'teacher', 'instructor', 'curriculum', 'student', 'academic', 'school'],
            'engineering': ['engineering', 'design', 'manufacturing', 'construction', 'technical', 'mechanical', 'electrical'],
            'legal': ['legal', 'law', 'attorney', 'lawyer', 'court', 'litigation', 'compliance', 'contract'],
            'creative': ['design', 'creative', 'art', 'graphic', 'visual', 'multimedia', 'photography', 'creative'],
            'hr': ['human resources', 'hr', 'recruitment', 'talent', 'employee', 'personnel', 'training'],
            'operations': ['operations', 'logistics', 'supply chain', 'procurement', 'vendor', 'process']
        }
        
        # Enhanced section keywords
        self.section_keywords = {
            'contact': ['contact', 'personal information', 'personal details', 'reach me', 'get in touch'],
            'summary': ['summary', 'objective', 'profile', 'about', 'overview', 'professional summary', 'career objective'],
            'education': ['education', 'academic', 'qualification', 'university', 'college', 'degree', 'certification', 'training'],
            'experience': ['experience', 'work history', 'employment', 'professional experience', 'career', 'work experience', 'professional background'],
            'skills': ['skills', 'technical skills', 'competencies', 'technologies', 'expertise', 'proficiencies', 'core competencies'],
            'projects': ['projects', 'portfolio', 'personal projects', 'key projects', 'notable projects'],
            'achievements': ['achievements', 'awards', 'honors', 'accomplishments', 'recognitions', 'accolades'],
            'certifications': ['certifications', 'certificates', 'licenses', 'professional certifications', 'credentials'],
            'languages': ['languages', 'linguistic skills', 'multilingual', 'language proficiency'],
            'interests': ['interests', 'hobbies', 'activities', 'personal interests', 'extracurricular'],
            'references': ['references', 'recommendations', 'referees', 'professional references']
        }
        
        # Industry-specific requirements
        self.industry_requirements = {
            'technology': {
                'essential_sections': ['skills', 'experience', 'projects'],
                'important_skills': ['programming', 'web_development', 'databases'],
                'preferred_length': (500, 900),
                'ats_weight': 0.4
            },
            'healthcare': {
                'essential_sections': ['education', 'experience', 'certifications'],
                'important_skills': ['medical_skills', 'healthcare_tech'],
                'preferred_length': (600, 1000),
                'ats_weight': 0.2
            },
            'finance': {
                'essential_sections': ['education', 'experience', 'certifications'],
                'important_skills': ['finance', 'accounting'],
                'preferred_length': (500, 800),
                'ats_weight': 0.35
            },
            'creative': {
                'essential_sections': ['projects', 'experience', 'skills'],
                'important_skills': ['graphic_design', 'web_design', 'multimedia'],
                'preferred_length': (400, 700),
                'ats_weight': 0.1
            },
            'marketing': {
                'essential_sections': ['experience', 'skills', 'achievements'],
                'important_skills': ['digital_marketing', 'traditional_marketing', 'sales'],
                'preferred_length': (500, 800),
                'ats_weight': 0.3
            },
            'education': {
                'essential_sections': ['education', 'experience', 'certifications'],
                'important_skills': ['teaching', 'educational_tech'],
                'preferred_length': (600, 900),
                'ats_weight': 0.25
            },
            'engineering': {
                'essential_sections': ['education', 'experience', 'skills'],
                'important_skills': ['mechanical_engineering', 'electrical_engineering', 'civil_engineering'],
                'preferred_length': (600, 1000),
                'ats_weight': 0.35
            },
            'legal': {
                'essential_sections': ['education', 'experience', 'certifications'],
                'important_skills': ['legal', 'regulatory'],
                'preferred_length': (700, 1200),
                'ats_weight': 0.2
            },
            'hr': {
                'essential_sections': ['education', 'experience', 'skills'],
                'important_skills': ['hr', 'hr_tech'],
                'preferred_length': (500, 800),
                'ats_weight': 0.3
            },
            'operations': {
                'essential_sections': ['experience', 'skills', 'education'],
                'important_skills': ['operations', 'supply_chain'],
                'preferred_length': (500, 800),
                'ats_weight': 0.35
            }
        }

    def detect_cv_type(self, text):
        """Automatically detect CV type based on content"""
        text_lower = text.lower()
        industry_scores = {}
        
        for industry, keywords in self.industry_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            industry_scores[industry] = score
        
        if not industry_scores or max(industry_scores.values()) == 0:
            return 'general'
        
        detected_industry = max(industry_scores, key=industry_scores.get)
        confidence = industry_scores[detected_industry] / sum(industry_scores.values())
        
        return {
            'primary_industry': detected_industry,
            'confidence': round(confidence, 2),
            'industry_scores': industry_scores
        }
        
    def extract_experience_duration(self, text):
        """Extract years of experience from CV"""
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*years?\s*in',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?'
        ]
        
        years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(match) for match in matches])
        
        # Also try to extract from date ranges
        date_patterns = [
            r'(\d{4})\s*[-–]\s*(\d{4})',
            r'(\d{4})\s*[-–]\s*present',
            r'(\d{4})\s*[-–]\s*current'
        ]
        
        current_year = datetime.now().year
        total_experience = 0
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                start_year = int(match[0])
                end_year = current_year if match[1] in ['present', 'current'] else int(match[1])
                total_experience += max(0, end_year - start_year)
        
        if years:
            explicit_years = max(years)
        else:
            explicit_years = 0
            
        return {
            'explicit_years': explicit_years,
            'calculated_years': total_experience,
            'estimated_years': max(explicit_years, total_experience)
        }
        
    def extract_education_level(self, text):
        """Extract education level from CV"""
        education_levels = {
            'phd': ['ph.d', 'phd', 'doctorate', 'doctoral'],
            'masters': ['master', 'mba', 'ms', 'ma', 'msc', 'm.s', 'm.a'],
            'bachelors': ['bachelor', 'bs', 'ba', 'bsc', 'b.s', 'b.a', 'undergraduate'],
            'associates': ['associate', 'aa', 'as', 'a.a', 'a.s'],
            'diploma': ['diploma', 'certificate', 'certification'],
            'high_school': ['high school', 'secondary', 'matriculation']
        }
        
        text_lower = text.lower()
        found_levels = []
        
        for level, keywords in education_levels.items():
            if any(keyword in text_lower for keyword in keywords):
                found_levels.append(level)
        
        # Return highest level found
        level_hierarchy = ['phd', 'masters', 'bachelors', 'associates', 'diploma', 'high_school']
        for level in level_hierarchy:
            if level in found_levels:
                return level
        
        return 'not_specified'
        
    def extract_certifications(self, text):
        """Extract professional certifications"""
        cert_patterns = [
            # Technology
            r'\b(?:aws|azure|gcp|google cloud)\s*certified\b',
            r'\b(?:cissp|cism|ceh|comptia)\b',
            r'\b(?:pmp|prince2|scrum master|csm)\b',
            # Healthcare
            r'\b(?:bls|acls|pals|cpr)\s*certified\b',
            r'\b(?:rn|lpn|cna|np)\b',
            # Finance
            r'\b(?:cpa|cfa|frm|cfp)\b',
            # General
            r'\bcertified\s+[\w\s]+\b',
            r'\b[\w\s]+\s+certification\b'
        ]
        
        certifications = []
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        return list(set(certifications))
        
    def analyze_keyword_density(self, text, job_description=None):
        """Analyze keyword density and relevance"""
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        word_count = Counter(words)
        total_words = len(words)
        
        # Calculate keyword density for skills
        skill_density = {}
        for category, skills in self.skill_keywords.items():
            category_count = 0
            for skill in skills:
                skill_words = skill.lower().split()
                if len(skill_words) == 1:
                    category_count += word_count.get(skill_words[0], 0)
                else:
                    # Multi-word skills
                    if skill.lower() in text.lower():
                        category_count += 1
            
            if total_words > 0:
                skill_density[category] = round((category_count / total_words) * 100, 2)
            else:
                skill_density[category] = 0
        
        return {
            'total_unique_words': len(word_count),
            'total_words': total_words,
            'skill_density': skill_density,
            'top_keywords': dict(word_count.most_common(20))
        }
        
    def check_grammar_and_readability(self, text):
        """Basic grammar and readability analysis"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Calculate simple readability scores
        words = text.split()
        total_words = len(words)
        total_sentences = len(sentences)
        
        # Simple flesch reading ease approximation
        if total_sentences > 0 and total_words > 0:
            avg_sentence_length = total_words / total_sentences
            # Count syllables (simple approximation)
            syllables = sum(self.count_syllables(word) for word in words)
            avg_syllables_per_word = syllables / total_words if total_words > 0 else 0
            
            # Simplified Flesch formula
            flesch_score = 206.835 - (1.015 * avg_sentence_length) - (84.6 * avg_syllables_per_word)
            flesch_score = max(0, min(100, flesch_score))  # Clamp between 0-100
        else:
            flesch_score = 50  # Default neutral score
            avg_sentence_length = 0
        
        # Basic grammar checks
        grammar_issues = []
        
        # Check for common issues
        if re.search(r'\bi\b', text):  # Lowercase 'i'
            grammar_issues.append("Use uppercase 'I' for first person")
        
        if len(re.findall(r'[.!?]', text)) < len(sentences) * 0.8:
            grammar_issues.append("Missing punctuation in some sentences")
        
        return {
            'flesch_reading_ease': round(flesch_score, 1),
            'readability_level': self.get_readability_level(flesch_score),
            'avg_sentence_length': round(avg_sentence_length, 1),
            'total_sentences': len(sentences),
            'grammar_suggestions': grammar_issues
        }
    
    def count_syllables(self, word):
        """Simple syllable counting approximation"""
        word = word.lower()
        vowels = 'aeiouy'
        syllable_count = 0
        prev_was_vowel = False
        
        for char in word:
            if char in vowels:
                if not prev_was_vowel:
                    syllable_count += 1
                prev_was_vowel = True
            else:
                prev_was_vowel = False
        
        # Handle silent 'e'
        if word.endswith('e') and syllable_count > 1:
            syllable_count -= 1
        
        return max(1, syllable_count)  # Every word has at least 1 syllable
        
    def get_readability_level(self, flesch_score):
        """Convert Flesch score to readability level"""
        if flesch_score >= 90:
            return "Very Easy"
        elif flesch_score >= 80:
            return "Easy"
        elif flesch_score >= 70:
            return "Fairly Easy"
        elif flesch_score >= 60:
            return "Standard"
        elif flesch_score >= 50:
            return "Fairly Difficult"
        elif flesch_score >= 30:
            return "Difficult"
        else:
            return "Very Difficult"
            
    def analyze_cv_completeness(self, sections, cv_type):
        """Analyze CV completeness based on industry standards"""
        completeness_score = 0
        missing_sections = []
        
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
        else:
            industry = cv_type if cv_type in self.industry_requirements else 'general'
        
        if industry in self.industry_requirements:
            requirements = self.industry_requirements[industry]
            essential_sections = requirements['essential_sections']
            
            for section in essential_sections:
                if section in sections and sections[section]:
                    completeness_score += 30
                else:
                    missing_sections.append(section)
        
        # Standard sections
        standard_sections = ['contact', 'summary', 'education', 'experience']
        for section in standard_sections:
            if section in sections and sections[section]:
                completeness_score += 10
            elif section not in missing_sections:
                missing_sections.append(section)
        
        return {
            'completeness_score': min(completeness_score, 100),
            'missing_sections': missing_sections,
            'industry_specific': industry != 'general'
        }
        
    def generate_industry_specific_feedback(self, cv_type, skills, sections, structure_info):
        """Generate industry-specific feedback"""
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
        else:
            industry = cv_type if cv_type in self.industry_requirements else 'general'
        
        feedback = []
        
        if industry == 'technology':
            if not any(skills.get(cat, []) for cat in ['programming', 'web_development', 'databases']):
                feedback.append("⚠ Tech CVs should highlight programming languages and technical skills")
            if 'projects' not in sections:
                feedback.append("⚠ Include a projects section to showcase your work")
            if structure_info['word_count'] > 1000:
                feedback.append("⚠ Tech CVs should be concise - consider reducing length")
                
        elif industry == 'healthcare':
            if not skills.get('medical_skills', []) and not skills.get('nursing', []):
                feedback.append("⚠ Healthcare CVs should emphasize clinical skills and experience")
            if 'certifications' not in sections:
                feedback.append("✗ Medical certifications are crucial for healthcare roles")
                
        elif industry == 'finance':
            if not any(skills.get(cat, []) for cat in ['finance', 'accounting']):
                feedback.append("⚠ Finance CVs should highlight analytical and financial skills")
            if 'certifications' not in sections:
                feedback.append("⚠ Professional certifications (CPA, CFA) are valuable in finance")
                
        elif industry == 'creative':
            if 'projects' not in sections:
                feedback.append("✗ Creative CVs must include a portfolio or projects section")
            if not any(skills.get(cat, []) for cat in ['graphic_design', 'web_design', 'multimedia']):
                feedback.append("⚠ Showcase your creative tools and software proficiency")
        
        return feedback
        
    # Enhanced extraction methods
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file with better error handling"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Error reading page {page_num}: {str(e)}")
                        continue
                return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file with table support"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def extract_text_from_txt(self, file_path):
        """Extract text from TXT file with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"Error reading TXT: {str(e)}"
        
        return "Error: Could not decode text file"
    
    def extract_text(self, file_path):
        """Extract text based on file extension"""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            return self.extract_text_from_txt(file_path)
        else:
            return "Unsupported file format. Please use PDF, DOCX, or TXT files."
    
    def extract_contact_info(self, text):
        """Enhanced contact information extraction"""
        contact_info = {}
        
        # Email extraction (improved)
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['emails'] = list(set(emails))
        
        # Phone extraction (enhanced patterns)
        phone_patterns = [
            r'(\+\d{1,4}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}',
            r'\+\d{1,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}',
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}'
        ]
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        contact_info['phones'] = list(set(phones))
        
        # Social media and professional profiles
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/view\?id=)[\w-]+'
        linkedin = re.findall(linkedin_pattern, text.lower())
        contact_info['linkedin'] = linkedin
        
        github_pattern = r'(?:github\.com/|git\.io/)[\w-]+'
        github = re.findall(github_pattern, text.lower())
        contact_info['github'] = github
        
        twitter_pattern = r'(?:twitter\.com/|@)[\w-]+'
        twitter = re.findall(twitter_pattern, text.lower())
        contact_info['twitter'] = twitter
        
        # Location extraction
        location_patterns = [
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, State
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, Country
        ]
        locations = []
        for pattern in location_patterns:
            locations.extend(re.findall(pattern, text))
        contact_info['locations'] = locations
        
        # Website extraction (improved)
        website_pattern = r'(?:https?://)?(?:www\.)?[\w-]+\.[\w.-]+(?:/[\w.-]*)*'
        websites = re.findall(website_pattern, text.lower())
        contact_info['websites'] = [w for w in websites if not any(social in w for social in ['linkedin', 'github', 'twitter'])]
        
        return contact_info
    
    def extract_skills(self, text):
        """Enhanced skills extraction with industry context"""
        text_lower = text.lower()
        found_skills = {}
        
        for category, skills in self.skill_keywords.items():
            found_skills[category] = []
            for skill in skills:
                # Check for exact matches and variations
                skill_lower = skill.lower()
                if skill_lower in text_lower:
                    found_skills[category].append(skill)
                # Check for skill variations (e.g., "JavaScript" vs "JS")
                elif category == 'programming':
                    variations = {
                        'javascript': ['js', 'node'],
                        'python': ['py'],
                        'typescript': ['ts']
                    }
                    if skill_lower in variations:
                        for var in variations[skill_lower]:
                            if var in text_lower:
                                found_skills[category].append(skill)
                                break
        
        return found_skills
    
    def identify_sections(self, text):
        """Enhanced section identification"""
        sections = {}
        lines = text.split('\n')
        
        current_section = 'general'
        sections[current_section] = []
        
        for line in lines:
            line_lower = line.lower().strip()
            section_found = False
            
            if not line_lower or len(line_lower) < 3:
                continue
            
            # Check for section headers
            for section, keywords in self.section_keywords.items():
                if any(keyword in line_lower for keyword in keywords) and len(line.strip()) < 60:
                    current_section = section
                    if current_section not in sections:
                        sections[current_section] = []
                    section_found = True
                    break
            
            if not section_found:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        return sections
    
    def analyze_length_and_structure(self, text):
        """Enhanced structure analysis"""
        words = text.split()
        word_count = len(words)
        char_count = len(text)
        lines = text.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        line_count = len(non_empty_lines)
        
        # Enhanced bullet point detection
        bullet_patterns = [
            r'^\s*[-•*▪▫◦‣⁃]\s',
            r'^\s*\d+\.\s',
            r'^\s*[a-zA-Z]\.\s',
            r'^\s*[ivxlc]+\.\s'
        ]
        
        bullet_count = 0
        for line in lines:
            if any(re.match(pattern, line) for pattern in bullet_patterns):
                bullet_count += 1
        
        # Calculate various metrics
        avg_words_per_line = word_count / line_count if line_count > 0 else 0
        avg_chars_per_word = char_count / word_count if word_count > 0 else 0
        
        # Count formatting elements
        bold_count = len(re.findall(r'\*\*.*?\*\*', text))
        italic_count = len(re.findall(r'\*.*?\*', text))
        caps_words = len(re.findall(r'\b[A-Z]{2,}\b', text))
        
        return {
            'word_count': word_count,
            'character_count': char_count,
            'line_count': line_count,
            'bullet_points': bullet_count,
            'avg_words_per_line': round(avg_words_per_line, 1),
            'avg_chars_per_word': round(avg_chars_per_word, 1),
            'formatting_elements': {
                'bold_text': bold_count,
                'italic_text': italic_count,
                'caps_words': caps_words
            }
        }
    
    def score_contact_section(self, contact_info):
        """Enhanced contact scoring"""
        score = 0
        feedback = []
        max_score = 100
        
        if contact_info['emails']:
            score += 25
            feedback.append("✓ Email address provided")
            # Check for professional email
            if any('@gmail.com' not in email and '@yahoo.com' not in email 
                   for email in contact_info['emails']):
                score += 5
                feedback.append("✓ Professional email domain")
        else:
            feedback.append("✗ Missing email address - Essential for contact")
        
        if contact_info['phones']:
            score += 20
            feedback.append("✓ Phone number provided")
        else:
            feedback.append("✗ Missing phone number - Important for contact")
        
        if contact_info['linkedin']:
            score += 20
            feedback.append("✓ LinkedIn profile included")
        else:
            feedback.append("⚠ Consider adding LinkedIn profile")
        
        if contact_info['github']:
            score += 15
            feedback.append("✓ GitHub profile included")
        else:
            feedback.append("⚠ Consider adding GitHub profile (for technical roles)")
        
        if contact_info['locations']:
            score += 10
            feedback.append("✓ Location information provided")
        
        if contact_info['websites']:
            score += 5
            feedback.append("✓ Personal website/portfolio included")
        
        if contact_info.get('twitter'):
            score += 5
            feedback.append("✓ Additional social media presence")
        
        return min(score, max_score), feedback
    
    def score_skills_section(self, skills, cv_type):
        """Industry-aware skills scoring"""
        score = 0
        feedback = []
        max_score = 100
        
        total_skills = sum(len(skill_list) for skill_list in skills.values())
        
        # Base scoring
        if total_skills >= 20:
            score += 40
            feedback.append(f"✓ Excellent variety of skills ({total_skills} skills found)")
        elif total_skills >= 15:
            score += 35
            feedback.append(f"✓ Very good variety of skills ({total_skills} skills found)")
        elif total_skills >= 10:
            score += 25
            feedback.append(f"✓ Good variety of skills ({total_skills} skills found)")
        elif total_skills >= 5:
            score += 15
            feedback.append(f"⚠ Moderate skills listed ({total_skills} skills found)")
        else:
            score += 5
            feedback.append(f"✗ Limited skills listed ({total_skills} skills found)")
        
        # Industry-specific scoring
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            if industry in self.industry_requirements:
                important_skills = self.industry_requirements[industry]['important_skills']
                industry_skill_count = sum(len(skills.get(cat, [])) for cat in important_skills)
                
                if industry_skill_count >= 5:
                    score += 25
                    feedback.append(f"✓ Strong {industry} industry skills")
                elif industry_skill_count >= 3:
                    score += 15
                    feedback.append(f"✓ Good {industry} industry skills")
                else:
                    feedback.append(f"⚠ Consider adding more {industry}-specific skills")
        
        # Skill diversity
        categories_with_skills = sum(1 for skill_list in skills.values() if skill_list)
        if categories_with_skills >= 6:
            score += 20
            feedback.append("✓ Excellent balance across skill categories")
        elif categories_with_skills >= 4:
            score += 15
            feedback.append("✓ Good balance across skill categories")
        elif categories_with_skills >= 2:
            score += 10
            feedback.append("⚠ Consider adding more diverse skills")
        
        # Soft skills
        if skills.get('leadership') or skills.get('interpersonal'):
            score += 10
            feedback.append("✓ Leadership and interpersonal skills included")
        
        if skills.get('project_management'):
            score += 5
            feedback.append("✓ Project management skills present")
        
        return min(score, max_score), feedback
    
    def score_structure_and_length(self, structure_info, cv_type):
        """Industry-aware structure scoring"""
        score = 0
        feedback = []
        max_score = 100
        
        word_count = structure_info['word_count']
        
        # Get industry-specific length preferences
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            if industry in self.industry_requirements:
                min_words, max_words = self.industry_requirements[industry]['preferred_length']
            else:
                min_words, max_words = 400, 800
        else:
            min_words, max_words = 400, 800
        
        # Length scoring
        if min_words <= word_count <= max_words:
            score += 50
            feedback.append(f"✓ Excellent length for your industry ({word_count} words)")
        elif (min_words - 100) <= word_count < min_words:
            score += 40
            feedback.append(f"✓ Good length ({word_count} words)")
        elif max_words < word_count <= (max_words + 200):
            score += 40
            feedback.append(f"✓ Acceptable length ({word_count} words)")
        elif word_count < (min_words - 100):
            score += 20
            feedback.append(f"⚠ Too short for your industry ({word_count} words)")
        else:
            score += 15
            feedback.append(f"⚠ Too long ({word_count} words)")
        
        # Structure elements
        bullet_points = structure_info['bullet_points']
        if bullet_points >= 8:
            score += 25
            feedback.append(f"✓ Excellent use of bullet points ({bullet_points})")
        elif bullet_points >= 5:
            score += 20
            feedback.append(f"✓ Good use of bullet points ({bullet_points})")
        elif bullet_points >= 2:
            score += 10
            feedback.append(f"✓ Some bullet points used ({bullet_points})")
        else:
            score += 5
            feedback.append("⚠ Consider using more bullet points for readability")
        
        # Line structure
        avg_words_per_line = structure_info['avg_words_per_line']
        if 6 <= avg_words_per_line <= 15:
            score += 15
            feedback.append("✓ Good line structure and readability")
        else:
            score += 8
            feedback.append("⚠ Consider optimizing line length for readability")
        
        # Formatting elements
        formatting = structure_info['formatting_elements']
        if formatting['bold_text'] > 0 or formatting['caps_words'] > 0:
            score += 10
            feedback.append("✓ Good use of formatting for emphasis")
        
        return min(score, max_score), feedback
    
    def score_sections(self, sections, cv_type):
        """Industry-aware section scoring"""
        score = 0
        feedback = []
        max_score = 100
        
        # Get industry requirements
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            if industry in self.industry_requirements:
                essential_sections = self.industry_requirements[industry]['essential_sections']
            else:
                essential_sections = ['education', 'experience', 'skills']
        else:
            essential_sections = ['education', 'experience', 'skills']
        
        # Essential sections scoring
        for section in essential_sections:
            if section in sections and sections[section]:
                content = ' '.join(sections[section])
                if len(content.strip()) > 100:
                    score += 25
                    feedback.append(f"✓ {section.capitalize()} section present with excellent content")
                elif len(content.strip()) > 50:
                    score += 20
                    feedback.append(f"✓ {section.capitalize()} section present with good content")
                else:
                    score += 10
                    feedback.append(f"⚠ {section.capitalize()} section present but needs more detail")
            else:
                feedback.append(f"✗ Missing essential {section} section")
        
        # Important sections
        important_sections = ['summary', 'projects', 'achievements']
        for section in important_sections:
            if section in sections and sections[section]:
                score += 10
                feedback.append(f"✓ {section.capitalize()} section present")
        
        # Bonus sections
        bonus_sections = ['certifications', 'languages', 'interests']
        bonus_count = sum(1 for section in bonus_sections if section in sections and sections[section])
        if bonus_count >= 2:
            score += 10
            feedback.append(f"✓ Additional sections enhance profile ({bonus_count} bonus sections)")
        elif bonus_count == 1:
            score += 5
            feedback.append("✓ Additional section adds value")
        
        return min(score, max_score), feedback
    
    def analyze_content_quality(self, text, sections):
        """Enhanced content quality analysis"""
        # Quantifiable achievements
        numbers_pattern = r'\b\d+(?:\.\d+)?(?:%|k|K|million|M|billion|B|x|X|\+)?\b'
        numbers_found = len(re.findall(numbers_pattern, text))
        
        # Action verbs (expanded list)
        action_verbs = [
            'achieved', 'developed', 'implemented', 'led', 'managed', 'created', 
            'improved', 'increased', 'reduced', 'designed', 'built', 'optimized',
            'delivered', 'coordinated', 'supervised', 'analyzed', 'established',
            'streamlined', 'automated', 'enhanced', 'collaborated', 'initiated',
            'executed', 'facilitated', 'mentored', 'negotiated', 'resolved'
        ]
        action_verb_count = sum(1 for verb in action_verbs if verb.lower() in text.lower())
        
        # Impact keywords
        impact_keywords = [
            'results', 'success', 'efficiency', 'performance', 'growth', 
            'savings', 'revenue', 'productivity', 'quality', 'innovation',
            'transformation', 'optimization', 'achievement', 'improvement',
            'solution', 'impact', 'breakthrough', 'milestone'
        ]
        impact_count = sum(1 for keyword in impact_keywords if keyword.lower() in text.lower())
        
        # Professional language
        professional_words = [
            'strategic', 'analytical', 'comprehensive', 'systematic', 'innovative',
            'collaborative', 'proactive', 'efficient', 'effective', 'dynamic'
        ]
        professional_count = sum(1 for word in professional_words if word.lower() in text.lower())
        
        quality_analysis = {
            'quantifiable_achievements': numbers_found,
            'action_verbs_used': action_verb_count,
            'impact_keywords': impact_count,
            'professional_language': professional_count,
            'has_summary': 'summary' in sections or 'objective' in sections,
            'has_achievements': 'achievements' in sections,
            'content_depth_score': min(100, (numbers_found * 5) + (action_verb_count * 3) + (impact_count * 4))
        }
        
        return quality_analysis
    
    def analyze_ats_compatibility(self, text, structure_info, cv_type):
        """Enhanced ATS compatibility analysis"""
        ats_score = 0
        ats_feedback = []
        
        # Get industry ATS weight
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            if industry in self.industry_requirements:
                ats_weight = self.industry_requirements[industry]['ats_weight']
            else:
                ats_weight = 0.3
        else:
            ats_weight = 0.3
        
        # Standard headers
        standard_headers = ['experience', 'education', 'skills', 'summary', 'contact']
        headers_found = sum(1 for header in standard_headers if header in text.lower())
        
        if headers_found >= 4:
            ats_score += 25
            ats_feedback.append("✓ Excellent use of standard section headers")
        elif headers_found >= 3:
            ats_score += 20
            ats_feedback.append("✓ Good use of standard section headers")
        else:
            ats_feedback.append("⚠ Use more standard section headers")
        
        # Clean formatting
        special_chars = len(re.findall(r'[^\w\s\-\.\,\(\)\@\:\/\%\&\#]', text))
        if special_chars < 30:
            ats_score += 20
            ats_feedback.append("✓ Clean formatting for ATS parsing")
        elif special_chars < 60:
            ats_score += 15
            ats_feedback.append("✓ Mostly clean formatting")
        else:
            ats_feedback.append("⚠ Reduce special characters for better ATS compatibility")
        
        # Contact information
        if '@' in text:
            ats_score += 15
            ats_feedback.append("✓ Email format is ATS-friendly")
        
        # File format compatibility
        ats_score += 10
        ats_feedback.append("✓ Standard file format (PDF/DOCX) used")
        
        # Keyword density
        total_words = structure_info['word_count']
        if total_words > 0:
            # Simple keyword density check
            important_word_count = len(re.findall(r'\b(?:experience|skill|manage|develop|lead|project)\b', text.lower()))
            keyword_density = (important_word_count / total_words) * 100
            
            if keyword_density >= 2:
                ats_score += 15
                ats_feedback.append("✓ Good keyword density for ATS systems")
            elif keyword_density >= 1:
                ats_score += 10
                ats_feedback.append("✓ Adequate keyword presence")
        
        # Line length
        avg_line_length = structure_info['character_count'] / structure_info['line_count']
        if avg_line_length < 80:
            ats_score += 15
            ats_feedback.append("✓ Good line length for ATS parsing")
        else:
            ats_feedback.append("⚠ Consider shorter lines for better parsing")
        
        # Apply industry weight
        weighted_score = ats_score * (ats_weight + 0.7)
        
        return min(int(weighted_score), 100), ats_feedback
    
    def calculate_overall_score(self, contact_score, skills_score, structure_score, sections_score, cv_type):
        """Industry-aware overall scoring"""
        # Base weights
        weights = {
            'contact': 0.15,
            'skills': 0.30,
            'structure': 0.25,
            'sections': 0.30
        }
        
        # Adjust weights based on industry
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            if industry == 'technology':
                weights['skills'] = 0.35
                weights['sections'] = 0.25
            elif industry == 'creative':
                weights['sections'] = 0.35  # Projects are crucial
                weights['skills'] = 0.25
            elif industry == 'healthcare':
                weights['sections'] = 0.35  # Certifications matter
                weights['contact'] = 0.10
        
        overall_score = (
            contact_score * weights['contact'] +
            skills_score * weights['skills'] +
            structure_score * weights['structure'] +
            sections_score * weights['sections']
        )
        
        return round(overall_score, 1)
    
    def generate_improvement_suggestions(self, contact_score, skills_score, structure_score, 
                                       sections_score, content_quality, cv_type, completeness):
        """Enhanced improvement suggestions"""
        suggestions = []
        
        # Industry context
        if isinstance(cv_type, dict):
            industry = cv_type['primary_industry']
            confidence = cv_type['confidence']
            
            if confidence < 0.7:
                suggestions.append({
                    'priority': 'Medium',
                    'area': 'Industry Focus',
                    'action': f'Strengthen {industry}-specific keywords and content',
                    'impact': 'Improves industry relevance and ATS matching'
                })
        
        # Critical Issues (Score < 50)
        if contact_score < 50:
            suggestions.append({
                'priority': 'Critical',
                'area': 'Contact Information',
                'action': 'Include complete professional contact details (email, phone, LinkedIn)',
                'impact': 'Essential for employer communication'
            })
        
        if sections_score < 50:
            missing_sections = completeness.get('missing_sections', [])
            suggestions.append({
                'priority': 'Critical',
                'area': 'Essential CV Sections',
                'action': f'Add missing sections: {", ".join(missing_sections)}',
                'impact': 'Required for CV completeness and professionalism'
            })
        
        # High Impact Improvements
        if skills_score < 70:
            suggestions.append({
                'priority': 'High',
                'area': 'Skills Portfolio',
                'action': 'Expand technical and soft skills with industry-relevant keywords',
                'impact': 'Significantly improves keyword matching and competency demonstration'
            })
        
        if content_quality['quantifiable_achievements'] < 5:
            suggestions.append({
                'priority': 'High',
                'area': 'Quantifiable Results',
                'action': 'Include specific metrics, percentages, and measurable outcomes',
                'impact': 'Demonstrates concrete professional value and impact'
            })
        
        if content_quality['action_verbs_used'] < 8:
            suggestions.append({
                'priority': 'High',
                'area': 'Dynamic Language',
                'action': 'Use more action verbs (achieved, developed, led, optimized)',
                'impact': 'Creates more engaging and impactful descriptions'
            })
        
        # Medium Impact Improvements
        if structure_score < 80:
            suggestions.append({
                'priority': 'Medium',
                'area': 'Document Structure',
                'action': 'Optimize word count, bullet points, and visual hierarchy',
                'impact': 'Enhances professional presentation and readability'
            })
        
        if not content_quality['has_summary']:
            suggestions.append({
                'priority': 'Medium',
                'area': 'Professional Summary',
                'action': 'Add a compelling 3-4 line professional summary at the top',
                'impact': 'Provides strong first impression and context'
            })
        
        if content_quality['professional_language'] < 5:
            suggestions.append({
                'priority': 'Medium',
                'area': 'Professional Language',
                'action': 'Incorporate more strategic and analytical terminology',
                'impact': 'Elevates professional tone and industry credibility'
            })
        
        # Low Impact but Valuable
        if not content_quality['has_achievements']:
            suggestions.append({
                'priority': 'Low',
                'area': 'Recognition Section',
                'action': 'Consider adding achievements, awards, or recognition section',
                'impact': 'Differentiates you from other candidates'
            })
        
        return suggestions
    
    def analyze_job_match(self, cv_text, job_description):
        """Analyze how well CV matches job description"""
        cv_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', cv_text.lower()))
        job_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', job_description.lower()))
        
        # Calculate overlap
        common_words = cv_words.intersection(job_words)
        match_percentage = (len(common_words) / len(job_words)) * 100 if job_words else 0
        
        # Extract requirements from job description
        job_skills = []
        for category, skills in self.skill_keywords.items():
            for skill in skills:
                if skill.lower() in job_description.lower():
                    job_skills.append(skill)
        
        # Check CV coverage of job skills
        cv_skills = self.extract_skills(cv_text)
        cv_skill_list = []
        for skills_list in cv_skills.values():
            cv_skill_list.extend([skill.lower() for skill in skills_list])
        
        matching_skills = [skill for skill in job_skills if skill.lower() in cv_skill_list]
        missing_skills = [skill for skill in job_skills if skill.lower() not in cv_skill_list]
        
        skill_match_percentage = (len(matching_skills) / len(job_skills)) * 100 if job_skills else 0
        
        # Generate recommendations
        recommendations = []
        if skill_match_percentage < 70:
            recommendations.append({
                'priority': 'High',
                'action': f'Add missing skills: {", ".join(missing_skills[:5])}',
                'impact': 'Improves keyword matching for this role'
            })
        
        if match_percentage < 50:
            recommendations.append({
                'priority': 'Medium',
                'action': 'Incorporate more keywords from the job description',
                'impact': 'Increases ATS compatibility and relevance'
            })
        
        return {
            'overall_match_percentage': round(match_percentage, 1),
            'skill_match_percentage': round(skill_match_percentage, 1),
            'matching_skills': matching_skills,
            'missing_skills': missing_skills,
            'total_job_skills_found': len(job_skills),
            'recommendations': recommendations,
            'match_grade': 'Excellent' if match_percentage >= 80 else 
                         'Good' if match_percentage >= 60 else 
                         'Fair' if match_percentage >= 40 else 'Poor'
        }
    
    def analyze_cv(self, file_path):
        """Main enhanced CV analysis method"""
        try:
            # Extract text
            text = self.extract_text(file_path)
            if not text or "Error reading" in text or "Unsupported file format" in text:
                return {"error": "Could not extract text from file. Please ensure it's a valid PDF, DOCX, or TXT file."}
            
            if len(text.strip()) < 50:
                return {"error": "File appears to be empty or contains too little text to analyze."}
            
            # Detect CV type/industry
            cv_type = self.detect_cv_type(text)
            
            # Extract comprehensive information
            contact_info = self.extract_contact_info(text)
            skills = self.extract_skills(text)
            sections = self.identify_sections(text)
            structure_info = self.analyze_length_and_structure(text)
            content_quality = self.analyze_content_quality(text, sections)
            experience_info = self.extract_experience_duration(text)
            education_level = self.extract_education_level(text)
            certifications = self.extract_certifications(text)
            keyword_analysis = self.analyze_keyword_density(text)
            readability = self.check_grammar_and_readability(text)
            completeness = self.analyze_cv_completeness(sections, cv_type)
            
            # Calculate scores
            contact_score, contact_feedback = self.score_contact_section(contact_info)
            skills_score, skills_feedback = self.score_skills_section(skills, cv_type)
            structure_score, structure_feedback = self.score_structure_and_length(structure_info, cv_type)
            sections_score, sections_feedback = self.score_sections(sections, cv_type)
            ats_score, ats_feedback = self.analyze_ats_compatibility(text, structure_info, cv_type)
            
            overall_score = self.calculate_overall_score(
                contact_score, skills_score, structure_score, sections_score, cv_type
            )
            
            # Generate suggestions
            suggestions = self.generate_improvement_suggestions(
                contact_score, skills_score, structure_score, sections_score, 
                content_quality, cv_type, completeness
            )
            
            # Generate industry-specific feedback
            industry_feedback = self.generate_industry_specific_feedback(
                cv_type, skills, sections, structure_info
            )
            
            # Compile comprehensive results
            results = {
                'overall_score': overall_score,
                'cv_type': cv_type,
                'scores': {
                    'contact': round(contact_score, 1),
                    'skills': round(skills_score, 1),
                    'structure': round(structure_score, 1),
                    'sections': round(sections_score, 1),
                    'ats_compatibility': round(ats_score, 1),
                    'completeness': round(completeness['completeness_score'], 1),
                    'overall': overall_score
                },
                'feedback': {
                    'contact': contact_feedback,
                    'skills': skills_feedback,
                    'structure': structure_feedback,
                    'sections': sections_feedback,
                    'ats_compatibility': ats_feedback,
                    'industry_specific': industry_feedback
                },
                'detailed_analysis': {
                    'content_quality': content_quality,
                    'skills_breakdown': skills,
                    'sections_content': list(sections.keys()),
                    'structure_metrics': structure_info,
                    'experience_analysis': experience_info,
                    'education_level': education_level,
                    'certifications_found': certifications,
                    'keyword_analysis': keyword_analysis,
                    'readability': readability,
                    'completeness_analysis': completeness
                },
                'suggestions': suggestions,
                'grade': self.get_grade(overall_score),
                'contact_info_extracted': contact_info
            }
            
            return results
            
        except Exception as e:
            return {"error": f"An error occurred during analysis: {str(e)}"}
    
    def get_grade(self, score):
        """Enhanced grading system"""
        if score >= 90:
            return {
                "grade": "Outstanding", 
                "message": "Exceptional CV that demonstrates professional excellence",
                "level": "A+"
            }
        elif score >= 85:
            return {
                "grade": "Excellent", 
                "message": "CV demonstrates exceptional professional presentation",
                "level": "A"
            }
        elif score >= 75:
            return {
                "grade": "Very Good", 
                "message": "Strong professional CV with solid fundamentals",
                "level": "B+"
            }
        elif score >= 65:
            return {
                "grade": "Good", 
                "message": "Competent CV structure with clear improvement potential",
                "level": "B"
            }
        elif score >= 50:
            return {
                "grade": "Satisfactory", 
                "message": "Basic CV framework with significant development opportunities",
                "level": "C"
            }
        else:
            return {
                "grade": "Needs Major Improvement", 
                "message": "CV requires substantial enhancements across multiple areas",
                "level": "D"
            }

# Initialize the enhanced analyzer
analyzer = EnhancedCVAnalyzer()

# Enhanced API Routes
@app.route('/')
def index():
    return jsonify({
        'service': 'Enhanced Multi-Industry CV Analysis API',
        'status': 'running',
        'version': '2.0.0',
        'features': [
            'Multi-industry CV analysis (Technology, Healthcare, Finance, Marketing, etc.)',
            'Automatic industry detection',
            'Advanced content quality assessment',
            'Enhanced ATS compatibility analysis',
            'Comprehensive skill extraction across all industries',
            'Experience duration analysis',
            'Education level detection',
            'Professional certification identification',
            'Readability and grammar analysis',
            'Industry-specific improvement suggestions'
        ],
        'supported_industries': [
            'Technology & IT', 'Healthcare & Medical', 'Finance & Accounting',
            'Marketing & Sales', 'Education & Training', 'Engineering & Manufacturing',
            'Creative & Design', 'Legal & Compliance', 'Human Resources',
            'Operations & Supply Chain', 'General/Cross-industry'
        ],
        'endpoints': {
            'analyze': {
                'method': 'POST',
                'url': '/analyze',
                'description': 'Upload CV file for comprehensive multi-industry analysis',
                'parameters': 'file (form-data)',
                'supported_formats': ['PDF', 'DOCX', 'TXT']
            },
            'analyze_with_job': {
                'method': 'POST',
                'url': '/analyze_with_job',
                'description': 'Analyze CV against specific job description',
                'parameters': 'file (form-data), job_description (text)'
            },
            'health': {
                'method': 'GET', 
                'url': '/health',
                'description': 'API health check'
            },
            'industries': {
                'method': 'GET',
                'url': '/industries',
                'description': 'Get supported industries and their requirements'
            }
        }
    })

@app.route('/analyze', methods=['POST'])
def analyze_cv():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Analyze the CV
            results = analyzer.analyze_cv(file_path)
            
            # Clean up the uploaded file
            os.remove(file_path)
            
            return jsonify(results)
            
        except Exception as e:
            # Clean up the uploaded file in case of error
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'error': f'Analysis failed: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file type. Please upload PDF, DOCX, or TXT files.'}), 400

@app.route('/analyze_with_job', methods=['POST'])
def analyze_cv_with_job():
    """Analyze CV against a specific job description"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    job_description = request.form.get('job_description', '')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not job_description.strip():
        return jsonify({'error': 'Job description is required'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Get basic CV analysis
            results = analyzer.analyze_cv(file_path)
            
            if 'error' not in results:
                # Add job matching analysis
                cv_text = analyzer.extract_text(file_path)
                job_match_analysis = analyzer.analyze_job_match(cv_text, job_description)
                results['job_match_analysis'] = job_match_analysis
            
            # Clean up the uploaded file
            os.remove(file_path)
            
            return jsonify(results)
            
        except Exception as e:
            # Clean up the uploaded file in case of error
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'error': f'Analysis failed: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file type. Please upload PDF, DOCX, or TXT files.'}), 400

@app.route('/industries', methods=['GET'])
def get_industries():
    """Get supported industries and their requirements"""
    return jsonify({
        'supported_industries': list(analyzer.industry_keywords.keys()),
        'industry_requirements': analyzer.industry_requirements,
        'skill_categories': list(analyzer.skill_keywords.keys())
    })

@app.route('/health')
def health_check():
    return jsonify({
        'status': 'healthy', 
        'message': 'Enhanced Multi-Industry CV Analysis API is running',
        'version': '2.0.0',
        'capabilities': [
            'Multi-industry support',
            'Automatic CV type detection',
            'Advanced content analysis',
            'Industry-specific feedback'
        ]
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
